from docstring_extractor import get_docstrings
from docx import Document


def docstring_handler(source='main.py', target='file'):
    content = extract(source=source)

    document = Document()
    document = content_to_docx(document, content)
    document.save(target + '.docx')


def extract(source):
    with open(source) as file:
        return get_docstrings(file)


def content_to_docx(document, content):
    document.add_heading(content['type'] + ': ' + content['name'], 0)
    if not content['docstring'] is None:
        document.add_paragraph(content['docstring_text'])
    document = docx_content_handler(content['content'], document, 1)
    return document


def docx_content_handler(content, document, level=1):
    for i in content:
        document.add_heading(i['type'] + ': ' + i['name'], level=level)
        if not i['docstring'] is None:
            actl_content = i['docstring_text']
            par = actl_content.find(':param')
            p = document.add_paragraph("")
            if par >= 0:
                p.add_run(actl_content[:par])
                par2 = par + 6
                actl_content = actl_content[par2:]
                p.add_run("Parameter: ").bold = True
                par3 = actl_content.find(':')
                p.add_run(actl_content[:par3])
                par3 += 1
                actl_content = actl_content[par3:]
            ret = actl_content.find(':return:')
            if ret >= 0:
                p.add_run(actl_content[:ret])
                ret2 = ret + 8
                p.add_run("Return: ").bold = True
                p.add_run(actl_content[ret2:])
                actl_content = ''
            p.add_run(actl_content)
        if len(i['content']) > 0:
            if level == 5:
                level = 4
            document = docx_content_handler(i['content'], document, level + 1)
    return document
